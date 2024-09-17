import os
from openpyxl import load_workbook
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import replicate
from langchain_community.llms import Ollama
import datetime


def cutPretext(result):
    full_response = ""
    temp = ""
    for item in result:
        temp+=item
        if temp != 'Output Started:' or temp != 'Output Ended:':
            full_response += item
        if len(temp) == 4:
            temp = ""
    print(full_response)
    return full_response

def summarize(sentence_input):
    llm = Ollama(model="llama3")
    prompt = "You are a third person past tense generator and the pronouns you use to describe the third person is always they. You just convert sentences given to you from first person to third person past tense and just give the output.Always start with Output Started and always end with output ended. The sentences are here : "+sentence_input
    result = llm.invoke(prompt)
    return cutPretext(result)


def desiredOutcome(sentence_input):

    llm = Ollama(model="llama3")
    prompt = "Give the summary of the expected desired outcome by student after analysing the sentences given to you. Always start with Output Started and always end with output ended. The sentences are given here : " + sentence_input
    result = llm.invoke(prompt)
    return cutPretext(result)


def create_borders(table):
    for cell in table._cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        
        # Create border elements
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            if tc == 4:
                border.set(qn('w:sz'), '4')  # Border size
            border.set(qn('w:sz'), '5')  # Border size
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Border color
            tcBorders.append(border)
        
        tcPr.append(tcBorders)

def create_word_document(data, filename):
    document = Document()
    
    details = ["Student Name: ",'Student ID: ','Program Code: ', 'Status: ']
    answer_details=[data[4],data[7],data[10],data[14]]
    for i in range(len(details)):
        bold_paragraph = document.add_paragraph()
        bold_run = bold_paragraph.add_run(details[i])
        bold_run.bold = True
        print(answer_details[i])
        bold_paragraph.add_run(str(answer_details[i]))
    
    
    document.add_heading('CASE OWNERSHIP', level=1)
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    # table.rows
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text="Completed by:"
    hdr_cells[1].text = "Noel Soans"

    create_borders(table)

    document.add_heading('ISSUE/RESOLUTION', level=1)
    table = document.add_table(rows=2, cols=2)
    table.autofit = True

    
    # table.rows
    issue= ["Academic Issue", data[12],"Student’s desired outcome/resolution", desiredOutcome(data[13])]
    for i in range(2):
        hdr_cells = table.rows[i].cells
        hdr_cells[0].text=issue[i*2]
        hdr_cells[1].text=issue[i*2+1]

    create_borders(table)

    document.add_heading('RECORD OF SUPPORT', level=1)
    table = document.add_table(rows=5, cols=2)
    table.autofit = True

    for i in range(5):
        hdr_cells = table.rows[i].cells
        if i == 0:
            hdr_cells[0].text = 'DATE'
            hdr_cells[1].text = 'NOTES'
        elif i == 1:
            if isinstance(data[2], datetime.datetime):
                hdr_cells[0].text = data[2].strftime('%Y-%m-%d')  # Adjust the format as needed
                hdr_cells[1].text = "Case Information: "+ summarize(data[13])
            else:
                hdr_cells[0].text = str(data[2])
    create_borders(table)

    document.add_heading('CLOSED STATUS', level=1)
    table = document.add_table(rows=2, cols=2)
    table.autofit = True

    for i in range(2):
        hdr_cells = table.rows[i].cells
        if i == 0:
            hdr_cells[0].text = 'DATE CLOSED'
            hdr_cells[1].text = 'RESOLUTION OUTCOME'
    create_borders(table)

    document.save(filename)

def monitor_excel_file(excel_filename, word_filename):
    while True:
        wb = load_workbook(excel_filename)
        ws = wb.active
        max_row = ws.max_row
        print(max_row, len(ws[f"A{max_row}:S{max_row}"][0]), ws.cell(row=1, column=5).value, ws[1][0].value, f"A{max_row}:S{max_row}")
        if ws.cell(row=1, column=1).value is not None:
            row_data = [cell.value for cell in ws[f"A{1}:S{1}"][0]]
            print(row_data)
            create_word_document(row_data, word_filename)
            print("Word document created successfully.")
            wb.close()
            break
        else:
            wb.close()
            break

if __name__ == "__main__":
    print("Here")
    excel_filename = "your_excel_file.xlsx"
    word_filename = "Course Reflection.docx"
    monitor_excel_file(excel_filename, word_filename)


# full_response = ""
    # inverted_comma = False
    # count = 0
    # for item in output:
    #     if item == 'Output Started' or item == 'Output Ended':
    #         inverted_comma = True
    #         count+=1
    #         print(count)
    #     elif inverted_comma and count == 1:
    #         full_response += item
    #     full_response += item
    # print(full_response)
    # return full_response